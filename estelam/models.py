from django.db import models
import PyPDF2
import re
import pandas as pd
import docx


class File(models.Model):

    filename = models.CharField(max_length=200, null=True, blank=True)   
    
    pdf_file = models.FileField(upload_to="", null=True, blank=True)
    date_created = models.DateTimeField(auto_now_add=True)


    def convert(self):
        
        file_name = self.filename        
        doc = PyPDF2.PdfReader(self.pdf_file)
        name = file_name.split(".")[0]
        pages = len(doc.pages)
        search = 'ردیف  '
        final_text = []
        count = 0
        for page in range(pages):
            current_page = doc.pages[page]
            text = current_page.extract_text()
            final_text.append(text[410:-85])
            
            if re.findall(search, text):
                count_page = len(re.findall(search, text))
                count +=count_page

        #scrape data
        final_text = "".join(final_text)
        scrape = final_text.split("کاال      ضمیمه ")
        data = []
        num = 0
        
        for page in range(1, len(scrape)):
            num +=1
            mat_code = scrape[page][17:30]
            
            mat_info = scrape[page].split("ردیف")[1].split("شرح")[0].strip()
            
            entity = scrape[page].split("مقدار")[0].strip().split()[-1]
                    
            entity_val = scrape[page].split("واحد")[0].strip().split("کاال")[-1].strip()
            
            if "سازنده اصلی" in scrape[page]:
                ori_creater = scrape[page].split(":سازنده اصلی کاال ")[0].strip().split(":تایپ اصلی")[-1].strip()
                
                ori_type = scrape[page].split(":تایپ اصلی")[0].strip().split("مبلغ واحد :")[-1].strip()
            if "سازنده اصیل" in scrape[page]:
                
                ori_creater = scrape[page].split(":سازنده اصیل کاال ")[0].strip().split(":تایپ اصیل")[-1].strip()
                ori_type = scrape[page].split(":تایپ اصیل")[0].strip().split("مبلغ واحد :")[-1].strip()
                
            if "فلگ" in scrape[page]:
                flag = scrape[page].split("فلگ")[1].split("کاال ")[1].strip()
                if "سازنده اصیل کاال" in scrape[page]:
                    d_flag = scrape[page].split("فلگ")[0].split("سازنده اصیل کاال")[1].split("شرح")[0].strip()
                if "سازنده اصلی کاال" in scrape[page]:
                    d_flag = scrape[page].split("فلگ")[0].split("سازنده اصلی کاال")[1].split("شرح")[0].strip()
                
            if "لیست تایپ های پیشنهادی" in scrape[page]:
                code_p = scrape[page].split("لیست تایپ های پیشنهادی")[1].split("کد سازنده")[0].strip()
                name_p = scrape[page].split("لیست تایپ های پیشنهادی")[1].split("کد سازنده")[1].split("نام سازنده")[0].strip()
                type_p = " ".join(scrape[page].split("لیست تایپ های پیشنهادی")[1].split("کد سازنده")[1].split("نام سازنده")[1].split("تایپ پیشنهادی")[0].split())
            if "یلست تایپ های پیشنهادی" in scrape[page]:
                code_p = scrape[page].split("یلست تایپ های پیشنهادی")[1].split("کد سازنده")[0].strip()
                name_p = scrape[page].split("یلست تایپ های پیشنهادی")[1].split("کد سازنده")[1].split("نام سازنده")[0].strip()
                type_p = " ".join(scrape[page].split("یلست تایپ های پیشنهادی")[1].split("کد سازنده")[1].split("نام سازنده")[1].split("تایپ پیشنهادی")[0].split())


            if "سازنده اصلی" in scrape[page]:
                

                temp = {
                    "ردیف":num ,
                    "کد کالا":mat_code ,
                    "شرح کالا":mat_info ,
                    "مقدار":entity ,
                    "واحد":entity_val ,
                    "سازنده اصلی کالا":ori_creater ,
                    "تایپ اصلی":ori_type ,
                    }
                temp["فلگ"] = flag
                temp["شرح فلگ"] = d_flag
            elif "سازنده اصیل" in scrape[page]:
                print("INJA HAM")
                temp = {
                    "ردیف":num ,
                    "کد کالا":mat_code ,
                    "شرح کالا":mat_info ,
                    "مقدار":entity ,
                    "واحد":entity_val ,
                    "سازنده اصلی کالا":ori_creater ,
                    "تایپ اصلی":ori_type ,
                    }
                temp["فلگ"] = flag
                temp["شرح فلگ"] = d_flag
            else:
                temp = {
                    "ردیف":num ,
                    "کد کالا":mat_code ,
                    "شرح کالا":mat_info ,
                    "مقدار":entity ,
                    "واحد":entity_val ,
                    }
            
            
            
            if "لیست تایپ های پیشنهادی" in scrape[page]:
                temp["کد سازنده"] = code_p
                temp["نام سازنده"] = name_p
                temp["تایپ پیشنهادی"] = type_p
            if "یلست تایپ های پیشنهادی" in scrape[page]:
                temp["کد سازنده"] = code_p
                temp["نام سازنده"] = name_p
                temp["تایپ پیشنهادی"] = type_p
            data.append(temp)
            





        df = pd.DataFrame.from_dict(data)
        df.to_excel('static/files/%s.xlsx'%name)
        
        # Create an instance of a word document
        doc = docx.Document()
        
        # Add a Title to the document
        doc.add_heading('Tamin Sanat Hengam', 0)
        
        # Table data in a form of list
        table = doc.add_table(rows=1, cols=6)

            # Adding heading in the 1st row of the table
        row = table.rows[0].cells

        row[0].text = 'POS'
        row[1].text = 'MATERIAL CODE'
        row[2].text = 'UNIT'
        row[3].text = 'QTY'
        row[4].text = 'UNIT PRICE'
        row[5].text = 'TOTAL PRICE'  
        # Creating a table object
        n = 0
        for d in data:    
            obj = Materials()
            row = table.add_row().cells
            
            row[0].text = str(d["ردیف"])
            row[1].text = str(d["کد کالا"])

            obj.code = str(d["کد کالا"])
            obj.description = d["شرح کالا"]
            

            if d["واحد"] == "عدد":
                row[2].text = "Nu"
            row[3].text = d["مقدار"]
            row[4].text = ""
            row[5].text = ""
            
            row = table.add_row().cells
            row[0].text = "Description :"
            if "شرح فلگ" in d.keys():
                row[1].text = d["شرح فلگ"]
            else:
                row[1].text = "Tecnical Data"
            row[2].text = ""
            row[3].text = ""
            row[4].text = "TYPE"
            row[5].text = "Manufacturer"
            
            
            row = table.add_row().cells
            row[0].text = ""
            row[1].text = d["شرح کالا"]
            row[2].text = ""
            row[3].text = ""
            row[4].text = ""
            row[5].text = ""
            if "تایپ اصلی" in d.keys():
                obj.main_type = str(d["تایپ اصلی"])
                obj.creator = str(d["سازنده اصلی کالا"])
                row[4].text = d["تایپ اصلی"]
                row[5].text = d["سازنده اصلی کالا"]
            if "کد سازنده" in d.keys():
                obj.p_code = d["کد سازنده"]
                obj.p_name = d["نام سازنده"]
                obj.p_type = d["تایپ پیشنهادی"]
            #merge
            table.cell(2+3*n, 0).merge(table.cell(3+3*n,0))
            table.cell(2+3*n, 1).merge(table.cell(2+3*n,3))
            table.cell(3+3*n, 1).merge(table.cell(3+3*n,3))
            n +=1
            obj.save()

        doc.save('static/files/%snew.docx'%name)


    def save(self, *args, **kwargs):
        self.convert()
        super(File, self).save(*args, **kwargs)
    def __str__(self):
        return self.filename

class Materials(models.Model):

    code = models.CharField(max_length=200, null=True, blank=True)
    description = models.CharField(max_length=1000, null=True, blank=True)
    main_type = models.CharField(max_length=200, null=True, blank=True)
    creator = models.CharField(max_length=200, null=True, blank=True)
    p_code = models.CharField(max_length=200, null=True, blank=True)
    p_name = models.CharField(max_length=200, null=True, blank=True)
    p_type = models.CharField(max_length=200, null=True, blank=True)
    self_description = models.CharField(max_length=1000, null=True, blank=True)
    
    def __str__(self):
        return self.code


class AddCompany(models.Model):
    name = models.CharField(max_length=200, null=True, blank=True)   
    excel_file = models.FileField(null=True, blank=True)
    
    def convert(self):
        Companies.objects.all().delete()
        df = pd.read_excel(self.excel_file)
        brand = df[df.columns[0]].dropna()
        new = []
        for b in brand:
            if b not in new:
                new.append(b)
        for n in new:
            a = df[df[df.columns[0]] == n]
            obj = Companies()
            for index, row in a.iterrows():
                obj.brand = row[a.columns[0]]
                obj.parts = row[a.columns[1]]
                obj.phone = row[a.columns[2]]
                obj.fax = row[a.columns[3]]
                obj.address = row[a.columns[4]]
                obj.person = row[a.columns[5]]
                obj.mobile = row[a.columns[6]]
                obj.email = row[a.columns[7]]
                obj.company = row[a.columns[8]]
                obj.country = row[a.columns[9]]
                obj.save()

    def save(self, *args, **kwargs):
        self.convert()
        super(AddCompany, self).save(*args, **kwargs)
    def __str__(self):
        return self.name
    
class Companies(models.Model):
    brand = models.CharField(max_length=200, null=True, blank=True)   
    parts = models.CharField(max_length=200, null=True, blank=True) 
    phone = models.CharField(max_length=200, null=True, blank=True) 
    fax = models.CharField(max_length=200, null=True, blank=True) 
    address = models.CharField(max_length=200, null=True, blank=True) 
    person = models.CharField(max_length=200, null=True, blank=True) 
    mobile = models.CharField(max_length=200, null=True, blank=True) 
    email = models.CharField(max_length=200, null=True, blank=True) 
    company = models.CharField(max_length=200, null=True, blank=True) 
    country = models.CharField(max_length=200, null=True, blank=True) 
    def __str__(self):
        return self.brand